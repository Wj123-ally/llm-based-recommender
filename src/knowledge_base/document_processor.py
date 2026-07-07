"""
知识库文档处理器。

负责：
- 解析不同格式的上传文件（PDF、DOCX、TXT、MD）
- 将文档切分为适合检索的文本块
- 向量化并存入 Milvus 知识库专用 collection
- 删除文件时清理对应索引

设计原则：
- 索引失败不影响文件本身的保存（记录日志并向后兼容）
- 使用与商品检索相同的 embedding 模型，保持向量空间一致
"""

import logging
from pathlib import Path
from typing import Any

from langchain_core.documents import Document

from src.shared import (
    create_embedding_model,
    ensure_knowledge_collection,
    get_knowledge_collection_name,
)

logger = logging.getLogger(__name__)

# 文档分块大小（字符数）
CHUNK_SIZE = 500
# 分块之间的重叠字符数
CHUNK_OVERLAP = 50


def _get_embedding_model():
    """获取文本 embedding 模型（单例缓存在 shared 层）。"""
    return create_embedding_model()


def _parse_file(file_path: Path) -> str:
    """
    根据文件扩展名解析文件内容，返回纯文本。
    支持的格式：PDF、DOCX、TXT、MD。
    """
    extension = file_path.suffix.lower().lstrip(".")
    file_path_str = str(file_path)

    if extension == "txt" or extension == "md":
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                return file_path.read_text(encoding="gbk")
            except UnicodeDecodeError as exc:
                raise ValueError(
                    f"文件编码无法识别，请转换为 UTF-8：{file_path.name}"
                ) from exc

    if extension == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf：pip install pypdf>=3.0.0")

        reader = PdfReader(file_path_str)
        text_parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        if not text_parts:
            raise ValueError(f"PDF 文件没有可提取的文本：{file_path.name}")
        return "\n".join(text_parts)

    if extension == "docx":
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx：pip install python-docx>=0.8.0")

        doc = DocxDocument(file_path_str)
        text_parts: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        if not text_parts:
            raise ValueError(f"DOCX 文件没有可提取的文本：{file_path.name}")
        return "\n".join(text_parts)

    raise ValueError(f"不支持的文件类型：.{extension}")


def _chunk_text(text: str, file_id: str, filename: str) -> list[Document]:
    """
    将文本切分为固定大小的块，每条携带 file_id 和来源文件名。
    """
    try:
        from langchain_core.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )

    chunks = splitter.create_documents(
        texts=[text],
        metadatas=[
            {
                "file_id": file_id,
                "source_filename": filename,
                "chunk_size": CHUNK_SIZE,
            }
        ],
    )

    # 为每个 chunk 生成唯一 ID，格式：{file_id}_{chunk_index}
    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = index
        chunk.metadata["chunk_id"] = f"{file_id}_{index}"

    return chunks


def index_file(file_path: Path, file_id: str, filename: str) -> int:
    """
    解析文件、分块、向量化、写入知识库 Milvus collection。

    返回值：成功索引的文本块数量。

    异常：解析或向量化失败时抛出，调用方应捕获并记录日志。
    """
    try:
        text = _parse_file(file_path)
    except ImportError as exc:
        logger.warning("解析文件 %s 缺少依赖：%s", filename, exc)
        raise
    except Exception as exc:
        logger.exception("解析文件 %s 失败", filename)
        raise

    if not text.strip():
        raise ValueError(f"文件内容为空：{filename}")

    chunks = _chunk_text(text, file_id, filename)
    logger.info("文件 %s 解析出 %s 个文本块", filename, len(chunks))

    # 获取 embedding 模型并计算向量
    embedding_model = _get_embedding_model()

    # 获取第一个向量的维度，确保 collection 已创建
    first_vector = embedding_model.embed_documents([chunks[0].page_content])[0]
    ensure_knowledge_collection(len(first_vector))

    from src.retriever.milvus_store import upsert_documents

    # 分批写入，每批 20 条
    batch_size = 20
    for start in range(0, len(chunks), batch_size):
        batch_chunks = chunks[start : start + batch_size]
        batch_texts = [chunk.page_content for chunk in batch_chunks]
        batch_vectors = embedding_model.embed_documents(batch_texts)

        # 为每个 chunk 设置 doc id（Milvus 用 chunk_id 作为主键）
        for chunk in batch_chunks:
            chunk.id = chunk.metadata["chunk_id"]

        upsert_documents(
            get_knowledge_collection_name(),
            batch_chunks,
            batch_vectors,
        )

    logger.info("文件 %s 索引完成，共 %s 块", filename, len(chunks))
    return len(chunks)


def delete_file_index(file_id: str) -> int:
    """
    删除某个 file_id 在知识库 Milvus collection 中的所有 chunk。

    返回值：删除的文本块数量（可能为 0，表示该文件未被索引过）。

    实现方式：通过 Milvus 的 filter 表达式查询所有属于该 file_id 的 chunk，
    然后批量删除。
    """
    from src.retriever.milvus_store import get_milvus_client

    client = get_milvus_client()
    collection_name = get_knowledge_collection_name()

    if not client.has_collection(collection_name):
        logger.info("知识库 collection 不存在，无需清理")
        return 0

    try:
        # 查询所有属于该 file_id 的 chunk
        results = client.query(
            collection_name=collection_name,
            filter=f'file_id == "{file_id}"',
            output_fields=["id"],
        )
    except Exception:
        logger.exception("查询知识库 collection 失败")
        return 0

    if not results:
        logger.info("文件 %s 没有可清理的索引记录", file_id)
        return 0

    chunk_ids = [item["id"] for item in results]
    try:
        client.delete(collection_name=collection_name, pks=chunk_ids)
    except Exception:
        logger.exception("删除知识库索引失败")
        return 0

    logger.info("已从知识库索引中删除文件 %s 的 %s 个文本块", file_id, len(chunk_ids))
    return len(chunk_ids)


def get_knowledge_collection_stats() -> dict[str, Any]:
    """
    返回知识库 collection 的统计信息。
    用于管理界面或健康检查。
    """
    from src.retriever.milvus_store import count_collection

    try:
        count = count_collection(get_knowledge_collection_name())
        return {
            "collection_name": get_knowledge_collection_name(),
            "chunk_count": count,
        }
    except Exception:
        logger.exception("读取知识库统计信息失败")
        return {"collection_name": get_knowledge_collection_name(), "chunk_count": 0}
